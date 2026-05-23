"""
Generate the Word test report required by the course assignment.
"""

import csv
import json
import platform
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT_DIR / "output" / "doc"
RESULT_DIR = ROOT_DIR / "output" / "test-results"
REPORT_PATH = OUTPUT_DIR / "头盔检测系统测试报告.docx"
PERFORMANCE_CSV = RESULT_DIR / "performance_results.csv"
PERFORMANCE_CHART = RESULT_DIR / "performance_chart.png"
DATASET_SUMMARY_JSON = RESULT_DIR / "dataset_evaluation_summary.json"


FUNCTION_CASES = [
    ["F-01", "正常图片上传", "上传包含骑手/头盔的 JPG 图片", "返回检测统计、检测框列表和标注图", "通过"],
    ["F-02", "无目标图片", "上传纯色或无骑手图像", "接口成功返回，检测数量可为 0", "通过"],
    ["F-03", "模糊图片", "上传经过高斯模糊处理的样本", "系统不崩溃，返回可解释结果", "通过"],
    ["F-04", "暗光/强光图片", "上传亮度降低或提高的样本", "系统可完成推理并返回统计", "通过"],
    ["F-05", "非图片文件", "上传 text/plain 文件", "返回 400 错误", "通过"],
    ["F-06", "空图片文件", "上传 0 字节 image/jpeg", "返回 400 错误", "通过"],
    ["F-07", "批量混合上传", "同时上传合法图片和非法文件", "合法项成功，非法项记录失败", "通过"],
]

UNIT_CASES = [
    ["U-01", "模型文件缺失", "不存在的 .pt 路径", "抛出 FileNotFoundError", "通过"],
    ["U-02", "灰度图预处理", "二维灰度数组", "转换为三通道 BGR", "通过"],
    ["U-03", "空数组防御", "空 numpy 数组", "抛出 ValueError", "通过"],
    ["U-04", "检测结果格式", "Mock YOLO 输出", "返回 class/confidence/bbox/box_coords", "通过"],
    ["U-05", "类别统计", "Mock 检测结果", "classes 与 detection_count 正确", "通过"],
    ["U-06", "BGRA 编码", "四通道图片", "可转 Base64", "通过"],
]

DEFECTS = [
    ["D-01", "中", "上传目录依赖运行目录，换目录启动会保存到错误位置", "已修复", "改为 backend/uploads 绝对路径"],
    ["D-02", "中", "空文件和损坏图片错误信息不够明确", "已修复", "增加空文件与解码失败校验"],
    ["D-03", "低", "前端英文界面不便于课堂演示", "已修复", "改为中文演示页"],
    ["D-04", "低", "单元测试直接依赖真实模型，结果不稳定", "已修复", "核心单测使用 Mock 模型"],
]


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    return table


def read_performance_rows():
    if not PERFORMANCE_CSV.exists():
        return [["未运行", "-", "-", "-", "-", "运行 scripts/run_performance_test.py 后自动填充"]]

    with PERFORMANCE_CSV.open("r", encoding="utf-8", newline="") as file:
        reader = csv.DictReader(file)
        return [
            [
                row["variant"],
                row["width"],
                row["height"],
                row["avg_latency_ms"],
                row["fps"],
                row["detection_count"],
            ]
            for row in reader
        ]


def read_dataset_rows():
    if not DATASET_SUMMARY_JSON.exists():
        return [["未运行", "-", "-", "-", "-", "-", "运行 scripts/evaluate_dataset.py 后自动填充"]]

    summary = json.loads(DATASET_SUMMARY_JSON.read_text(encoding="utf-8"))
    rows = [
        [
            "Overall",
            summary["ground_truth_objects"],
            summary["predicted_objects"],
            summary["true_positive"],
            summary["precision"],
            summary["recall"],
            summary["f1"],
        ]
    ]
    for class_name, metrics in summary["per_class"].items():
        rows.append(
            [
                class_name,
                metrics["ground_truth"],
                metrics["predicted"],
                metrics["true_positive"],
                metrics["precision"],
                metrics["recall"],
                metrics["f1"],
            ]
        )
    return rows


def read_dataset_summary_text():
    if not DATASET_SUMMARY_JSON.exists():
        return "外部数据集测试尚未运行。"

    summary = json.loads(DATASET_SUMMARY_JSON.read_text(encoding="utf-8"))
    return (
        f"外部数据集测试使用 Kaggle andrewmvd/helmet-detection，共评估 "
        f"{summary['images_evaluated']} 张图片，{summary['ground_truth_objects']} 个标注目标；"
        f"置信度阈值 {summary['conf_threshold']}，IoU 匹配阈值 {summary['match_iou_threshold']}，"
        f"平均推理耗时 {summary['avg_latency_ms']} ms。"
    )


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_styles(document)

    title = document.add_heading("电动车骑手头盔检测系统测试报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("课程：软件测试与质量管理")
    document.add_paragraph("团队成员：姓名/学号待补充")
    document.add_paragraph("系统版本：头盔检测原型 v1.0")
    document.add_paragraph("报告说明：本报告依据作业要求整理功能测试、单元测试、性能测试、鲁棒性测试和缺陷分析。")

    document.add_heading("1. 测试环境", level=1)
    add_table(
        document,
        ["项目", "内容"],
        [
            ["操作系统", platform.platform()],
            ["Python", sys.version.split()[0]],
            ["后端框架", "FastAPI + OpenCV + Ultralytics YOLOv8"],
            ["前端框架", "React 18 + Vite"],
            ["模型文件", "backend/best.pt"],
            ["测试框架", "pytest, FastAPI TestClient"],
        ],
    )

    document.add_heading("2. 功能测试用例", level=1)
    add_table(document, ["编号", "场景", "输入/步骤", "预期结果", "结果"], FUNCTION_CASES)

    document.add_heading("3. 单元测试与覆盖说明", level=1)
    document.add_paragraph(
        "单元测试位于 backend/tests/test_detector.py，使用 Mock 模型隔离真实推理波动，重点覆盖预处理、后处理和异常输入。"
    )
    add_table(document, ["编号", "测试点", "输入", "预期结果", "结果"], UNIT_CASES)
    document.add_paragraph("覆盖率命令：python -m pytest backend/tests --cov=backend --cov-report=term-missing")

    document.add_heading("4. API 集成与鲁棒性测试", level=1)
    document.add_paragraph(
        "集成测试位于 backend/tests/test_api.py，覆盖 /health、单图上传、批量上传、阈值边界、非图片、空文件、损坏图片和检测器未初始化。"
    )
    add_table(
        document,
        ["类别", "覆盖内容", "判定标准"],
        [
            ["接口健康", "GET /health", "HTTP 200 且 status=healthy"],
            ["合法上传", "image/jpeg 文件", "HTTP 200 且返回检测统计"],
            ["异常输入", "text/plain、空文件、损坏图片", "HTTP 400 且错误信息明确"],
            ["服务异常", "detector=None", "HTTP 503"],
            ["边界参数", "conf_threshold 超出 0-1", "HTTP 422"],
        ],
    )

    document.add_heading("5. 性能测试", level=1)
    document.add_paragraph(
        "性能测试脚本 scripts/run_performance_test.py 使用真实 best.pt 模型，对不同分辨率和亮度/模糊样本记录平均耗时与 FPS。"
    )
    add_table(
        document,
        ["样本", "宽度", "高度", "平均耗时(ms)", "FPS", "检测数量"],
        read_performance_rows(),
    )
    if PERFORMANCE_CHART.exists():
        document.add_picture(str(PERFORMANCE_CHART), width=Inches(6.2))
    else:
        document.add_paragraph("性能柱状图将在运行性能脚本后生成。")

    document.add_heading("6. 外部数据集检测效果", level=1)
    document.add_paragraph(read_dataset_summary_text())
    add_table(
        document,
        ["类别", "标注目标", "预测目标", "正确匹配", "Precision", "Recall", "F1"],
        read_dataset_rows(),
    )

    document.add_heading("7. 缺陷列表", level=1)
    add_table(document, ["编号", "严重程度", "缺陷描述", "状态", "处理说明"], DEFECTS)

    document.add_heading("8. 结论", level=1)
    document.add_paragraph(
        "系统已满足课程作业中“基于图像的智能识别软件原型”要求，能够完成图片上传、头盔检测、结果标注和统计展示。"
        "自动化测试覆盖正常、异常、边界和鲁棒性场景；性能测试可为不同输入规模下的推理耗时提供数据。"
        "后续改进方向包括增加真实场景样本集、统计识别准确率、支持视频流检测和补充长期稳定性测试。"
    )

    document.save(REPORT_PATH)
    print(f"Wrote {REPORT_PATH}")


if __name__ == "__main__":
    main()
